# app.py - API Flask principale
from flask import Flask, request, jsonify, send_file
from flask_cors import CORS
from datetime import datetime, timedelta
import uuid
import os
from models import Devis, DevisItem
from pdf_generator import generate_pdf
from docx_generator import generate_docx

app = Flask(__name__)
CORS(app)

# Configuration
app.config['UPLOAD_FOLDER'] = 'generated'
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

@app.route('/health', methods=['GET'])
def health_check():
    return jsonify({"status": "healthy"}), 200

@app.route('/api/devis', methods=['POST'])
def create_devis():
    """
    Créer un nouveau devis avec les données reçues
    """
    try:
        data = request.json
        
        # Créer l'objet devis
        devis = Devis(
            numero=data.get('numero', f"D-{datetime.now().year}-{str(uuid.uuid4())[:3]}"),
            date_emission=data.get('date_emission', datetime.now().strftime('%d/%m/%Y')),
            date_expiration=data.get('date_expiration', (datetime.now() + timedelta(days=30)).strftime('%d/%m/%Y')),
            
            # Informations fournisseur
            fournisseur_nom=data.get('fournisseur_nom', 'Infinytia'),
            fournisseur_adresse=data.get('fournisseur_adresse', '61 Rue De Lyon'),
            fournisseur_ville=data.get('fournisseur_ville', '75012 Paris, FR'),
            fournisseur_email=data.get('fournisseur_email', 'contact@infinytia.com'),
            fournisseur_siret=data.get('fournisseur_siret', '93968736400017'),
            
            # Informations client
            client_nom=data.get('client_nom'),
            client_adresse=data.get('client_adresse'),
            client_ville=data.get('client_ville'),
            client_siret=data.get('client_siret'),
            client_tva=data.get('client_tva'),
            
            # Articles
            items=[]
        )
        
        # Ajouter les articles
        for item_data in data.get('items', []):
            item = DevisItem(
                description=item_data.get('description'),
                details=item_data.get('details', []),
                quantite=item_data.get('quantite', 1),
                prix_unitaire=item_data.get('prix_unitaire', 0),
                tva_taux=item_data.get('tva_taux', 20),
                remise=item_data.get('remise', 0)
            )
            devis.items.append(item)
        
        # Calculer les totaux
        devis.calculate_totals()
        
        # Format de sortie demandé
        output_format = data.get('format', 'pdf').lower()
        
        if output_format == 'pdf':
            filename = generate_pdf(devis)
            mimetype = 'application/pdf'
        elif output_format == 'docx':
            filename = generate_docx(devis)
            mimetype = 'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
        else:
            return jsonify({"error": "Format non supporté. Utilisez 'pdf' ou 'docx'"}), 400
        
        # Retourner le fichier
        return send_file(
            filename,
            mimetype=mimetype,
            as_attachment=True,
            download_name=f"devis_{devis.numero}.{output_format}"
        )
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

@app.route('/api/devis/preview', methods=['POST'])
def preview_devis():
    """
    Prévisualiser les données du devis sans générer de fichier
    """
    try:
        data = request.json
        
        # Créer l'objet devis
        devis = Devis(
            numero=data.get('numero', f"D-{datetime.now().year}-{str(uuid.uuid4())[:3]}"),
            date_emission=data.get('date_emission', datetime.now().strftime('%d/%m/%Y')),
            date_expiration=data.get('date_expiration', (datetime.now() + timedelta(days=30)).strftime('%d/%m/%Y')),
            
            # Informations fournisseur
            fournisseur_nom=data.get('fournisseur_nom', 'Infinytia'),
            fournisseur_adresse=data.get('fournisseur_adresse', '61 Rue De Lyon'),
            fournisseur_ville=data.get('fournisseur_ville', '75012 Paris, FR'),
            fournisseur_email=data.get('fournisseur_email', 'contact@infinytia.com'),
            fournisseur_siret=data.get('fournisseur_siret', '93968736400017'),
            
            # Informations client
            client_nom=data.get('client_nom'),
            client_adresse=data.get('client_adresse'),
            client_ville=data.get('client_ville'),
            client_siret=data.get('client_siret'),
            client_tva=data.get('client_tva'),
            
            # Articles
            items=[]
        )
        
        # Ajouter les articles
        for item_data in data.get('items', []):
            item = DevisItem(
                description=item_data.get('description'),
                details=item_data.get('details', []),
                quantite=item_data.get('quantite', 1),
                prix_unitaire=item_data.get('prix_unitaire', 0),
                tva_taux=item_data.get('tva_taux', 20),
                remise=item_data.get('remise', 0)
            )
            devis.items.append(item)
        
        # Calculer les totaux
        devis.calculate_totals()
        
        return jsonify(devis.to_dict()), 200
        
    except Exception as e:
        return jsonify({"error": str(e)}), 500

if __name__ == '__main__':
    port = int(os.environ.get('PORT', 5000))
    app.run(host='0.0.0.0', port=port)

# models.py - Modèles de données
from dataclasses import dataclass, field
from typing import List

@dataclass
class DevisItem:
    description: str
    details: List[str] = field(default_factory=list)
    quantite: int = 1
    prix_unitaire: float = 0
    tva_taux: float = 20
    remise: float = 0
    
    @property
    def total_ht(self):
        return (self.quantite * self.prix_unitaire) - self.remise
    
    @property
    def montant_tva(self):
        return self.total_ht * (self.tva_taux / 100)
    
    @property
    def total_ttc(self):
        return self.total_ht + self.montant_tva

@dataclass
class Devis:
    numero: str
    date_emission: str
    date_expiration: str
    
    # Fournisseur
    fournisseur_nom: str
    fournisseur_adresse: str
    fournisseur_ville: str
    fournisseur_email: str
    fournisseur_siret: str
    
    # Client
    client_nom: str
    client_adresse: str
    client_ville: str
    client_siret: str
    client_tva: str
    
    # Articles
    items: List[DevisItem] = field(default_factory=list)
    
    # Totaux
    total_ht: float = 0
    total_tva: float = 0
    total_ttc: float = 0
    
    def calculate_totals(self):
        self.total_ht = sum(item.total_ht for item in self.items)
        self.total_tva = sum(item.montant_tva for item in self.items)
        self.total_ttc = self.total_ht + self.total_tva
    
    def to_dict(self):
        return {
            "numero": self.numero,
            "date_emission": self.date_emission,
            "date_expiration": self.date_expiration,
            "fournisseur": {
                "nom": self.fournisseur_nom,
                "adresse": self.fournisseur_adresse,
                "ville": self.fournisseur_ville,
                "email": self.fournisseur_email,
                "siret": self.fournisseur_siret
            },
            "client": {
                "nom": self.client_nom,
                "adresse": self.client_adresse,
                "ville": self.client_ville,
                "siret": self.client_siret,
                "tva": self.client_tva
            },
            "items": [
                {
                    "description": item.description,
                    "details": item.details,
                    "quantite": item.quantite,
                    "prix_unitaire": item.prix_unitaire,
                    "tva_taux": item.tva_taux,
                    "remise": item.remise,
                    "total_ht": item.total_ht
                } for item in self.items
            ],
            "totaux": {
                "total_ht": self.total_ht,
                "total_tva": self.total_tva,
                "total_ttc": self.total_ttc
            }
        }

# pdf_generator.py - Génération PDF
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
import os

def generate_pdf(devis):
    filename = os.path.join('generated', f'devis_{devis.numero}.pdf')
    doc = SimpleDocTemplate(filename, pagesize=A4)
    elements = []
    styles = getSampleStyleSheet()
    
    # En-tête
    header_style = ParagraphStyle(
        'CustomHeader',
        parent=styles['Heading1'],
        fontSize=24,
        textColor=colors.HexColor('#333333'),
        spaceAfter=30
    )
    
    elements.append(Paragraph("Devis", header_style))
    
    # Informations du devis
    info_data = [
        ['Numéro de devis', devis.numero],
        ['Date d\'émission', devis.date_emission],
        ['Date d\'expiration', devis.date_expiration]
    ]
    
    info_table = Table(info_data, colWidths=[6*cm, 6*cm])
    info_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), colors.grey),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
    ]))
    
    elements.append(info_table)
    elements.append(Spacer(1, 20))
    
    # Fournisseur et Client
    company_data = [
        [Paragraph(f"<b>{devis.fournisseur_nom}</b>", styles['Normal']), 
         Paragraph(f"<b>{devis.client_nom}</b>", styles['Normal'])],
        [devis.fournisseur_adresse, devis.client_adresse],
        [devis.fournisseur_ville, devis.client_ville],
        [devis.fournisseur_email, f"SIRET: {devis.client_siret}"],
        [devis.fournisseur_siret, f"TVA: {devis.client_tva}"]
    ]
    
    company_table = Table(company_data, colWidths=[9*cm, 9*cm])
    company_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (-1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
    ]))
    
    elements.append(company_table)
    elements.append(Spacer(1, 30))
    
    # Articles
    items_data = [['Description', 'Qté', 'Prix unitaire', 'TVA (%)', 'Total HT']]
    
    for item in devis.items:
        # Ligne principale
        items_data.append([
            Paragraph(f"<b>{item.description}</b>", styles['Normal']),
            str(item.quantite),
            f"{item.prix_unitaire:.2f} €",
            f"{item.tva_taux} %",
            f"{item.total_ht:.2f} €"
        ])
        
        # Détails
        for detail in item.details:
            items_data.append([
                Paragraph(f"<font size='8'>{detail}</font>", styles['Normal']),
                '', '', '', ''
            ])
        
        # Remise si applicable
        if item.remise > 0:
            items_data.append([
                Paragraph("<font size='8' color='red'>Remise</font>", styles['Normal']),
                '', '', '', f"-{item.remise:.2f} €"
            ])
    
    items_table = Table(items_data, colWidths=[9*cm, 2*cm, 3*cm, 2*cm, 3*cm])
    items_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
        ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
        ('ALIGN', (1, 0), (-1, -1), 'RIGHT'),
        ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 9),
        ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
        ('GRID', (0, 0), (-1, -1), 1, colors.black),
    ]))
    
    elements.append(items_table)
    elements.append(Spacer(1, 20))
    
    # Totaux
    totals_data = [
        ['Total HT', f"{devis.total_ht:.2f} €"],
        ['Montant total de la TVA', f"{devis.total_tva:.2f} €"],
        ['Total TTC', f"{devis.total_ttc:.2f} €"]
    ]
    
    totals_table = Table(totals_data, colWidths=[10*cm, 4*cm])
    totals_table.setStyle(TableStyle([
        ('ALIGN', (0, 0), (-1, -1), 'RIGHT'),
        ('FONTNAME', (0, -1), (-1, -1), 'Helvetica-Bold'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BACKGROUND', (0, -1), (-1, -1), colors.grey),
        ('TEXTCOLOR', (0, -1), (-1, -1), colors.whitesmoke),
        ('LINEBELOW', (0, -2), (-1, -2), 1, colors.black),
    ]))
    
    elements.append(totals_table)
    
    # Générer le PDF
    doc.build(elements)
    return filename

# docx_generator.py - Génération DOCX
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os

def generate_docx(devis):
    filename = os.path.join('generated', f'devis_{devis.numero}.docx')
    doc = Document()
    
    # En-tête
    header = doc.add_heading('Devis', 0)
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Informations du devis
    doc.add_paragraph(f'Numéro de devis: {devis.numero}')
    doc.add_paragraph(f'Date d\'émission: {devis.date_emission}')
    doc.add_paragraph(f'Date d\'expiration: {devis.date_expiration}')
    doc.add_paragraph()
    
    # Tableau fournisseur/client
    table = doc.add_table(rows=5, cols=2)
    table.style = 'Light Grid'
    
    # Fournisseur
    table.cell(0, 0).text = devis.fournisseur_nom
    table.cell(1, 0).text = devis.fournisseur_adresse
    table.cell(2, 0).text = devis.fournisseur_ville
    table.cell(3, 0).text = devis.fournisseur_email
    table.cell(4, 0).text = devis.fournisseur_siret
    
    # Client
    table.cell(0, 1).text = devis.client_nom
    table.cell(1, 1).text = devis.client_adresse
    table.cell(2, 1).text = devis.client_ville
    table.cell(3, 1).text = f'SIRET: {devis.client_siret}'
    table.cell(4, 1).text = f'TVA: {devis.client_tva}'
    
    doc.add_paragraph()
    
    # Articles
    items_table = doc.add_table(rows=1, cols=5)
    items_table.style = 'Table Grid'
    
    # En-têtes
    headers = ['Description', 'Qté', 'Prix unitaire', 'TVA (%)', 'Total HT']
    for i, header in enumerate(headers):
        items_table.cell(0, i).text = header
    
    # Articles
    for item in devis.items:
        row = items_table.add_row()
        row.cells[0].text = item.description
        row.cells[1].text = str(item.quantite)
        row.cells[2].text = f'{item.prix_unitaire:.2f} €'
        row.cells[3].text = f'{item.tva_taux} %'
        row.cells[4].text = f'{item.total_ht:.2f} €'
        
        # Détails
        for detail in item.details:
            detail_row = items_table.add_row()
            detail_row.cells[0].text = f'  - {detail}'
        
        # Remise
        if item.remise > 0:
            remise_row = items_table.add_row()
            remise_row.cells[0].text = '  Remise'
            remise_row.cells[4].text = f'-{item.remise:.2f} €'
    
    doc.add_paragraph()
    
    # Totaux
    doc.add_paragraph(f'Total HT: {devis.total_ht:.2f} €')
    doc.add_paragraph(f'Montant total de la TVA: {devis.total_tva:.2f} €')
    doc.add_paragraph(f'Total TTC: {devis.total_ttc:.2f} €')
    
    # Sauvegarder
    doc.save(filename)
    return filename

# requirements.txt
Flask==3.0.0
flask-cors==4.0.0
reportlab==4.0.7
python-docx==1.1.0
gunicorn==21.2.0

# Procfile
web: gunicorn app:app

# railway.json
{
  "build": {
    "builder": "NIXPACKS"
  },
  "deploy": {
    "startCommand": "gunicorn app:app --bind 0.0.0.0:$PORT"
  }
}

# Exemple de requête n8n (HTTP Request)
{
  "method": "POST",
  "url": "https://your-app.railway.app/api/devis",
  "headers": {
    "Content-Type": "application/json"
  },
  "body": {
    "numero": "D-2025-061",
    "client_nom": "CRINSTALLE",
    "client_adresse": "400 AVENUE GEORGES CLEMENCEAU",
    "client_ville": "84300 CAVAILLON, FR",
    "client_siret": "952170736",
    "client_tva": "FR57952170736",
    "format": "pdf",
    "items": [
      {
        "description": "AGENT IA SLACK - GESTION DES ARMOIRES FIBRE",
        "details": [
          "Création de groupes Slack dédiés pour les équipes de techniciens",
          "Agent IA capable de retrouver les adresses des armoires à partir de leur référence",
          "Interface conversationnelle pour requêtes instantanées",
          "Base de données géolocalisée des armoires"
        ],
        "quantite": 1,
        "prix_unitaire": 900,
        "tva_taux": 20,
        "remise": 100
      },
      {
        "description": "SYSTÈME DE GESTION DES JETONS CLIENTS",
        "details": [
          "Agent IA récupérant les numéros de jetons",
          "Automatisation de l'envoi d'emails aux responsables concernés",
          "Traçabilité des demandes",
          "Interface de suivi en temps réel sur slack"
        ],
        "quantite": 1,
        "prix_unitaire": 850,
        "tva_taux": 20,
        "remise": 100
      }
    ]
  }
}
