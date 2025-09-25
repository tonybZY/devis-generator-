# docx_generator.py - Version améliorée avec support des factures, thèmes colorés et logos
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import os
import requests
from io import BytesIO

# Thèmes de couleurs pour DOCX (format RGB)
THEMES_COULEURS_DOCX = {
    'bleu': {
        'principale': RGBColor(44, 62, 80),
        'header_bg': "2d3436",
        'accent': RGBColor(52, 152, 219)
    },
    'vert': {
        'principale': RGBColor(39, 174, 96),
        'header_bg': "1e8449", 
        'accent': RGBColor(88, 214, 141)
    },
    'rouge': {
        'principale': RGBColor(231, 76, 60),
        'header_bg': "c0392b",
        'accent': RGBColor(241, 148, 138)
    },
    'violet': {
        'principale': RGBColor(155, 89, 182),
        'header_bg': "8e44ad",
        'accent': RGBColor(215, 189, 226)
    },
    'orange': {
        'principale': RGBColor(230, 126, 34),
        'header_bg': "d35400",
        'accent': RGBColor(245, 176, 65)
    },
    'noir': {
        'principale': RGBColor(44, 62, 80),
        'header_bg': "2c3e50",
        'accent': RGBColor(149, 165, 166)
    }
}

def set_cell_background(cell, color):
    """Définir la couleur de fond d'une cellule"""
    shading_elm = OxmlElement("w:shd")
    shading_elm.set(qn("w:fill"), color)
    cell._element.get_or_add_tcPr().append(shading_elm)

def download_and_add_logo(doc, logo_url):
    """Télécharger le logo et l'ajouter au document DOCX"""
    if not logo_url:
        return None
    
    try:
        # Télécharger l'image avec un timeout
        response = requests.get(logo_url, timeout=10)
        if response.status_code == 200:
            img_data = BytesIO(response.content)
            
            # Créer un paragraphe pour le logo aligné à droite
            logo_paragraph = doc.add_paragraph()
            logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            
            # Ajouter l'image avec une taille maximale
            run = logo_paragraph.runs[0] if logo_paragraph.runs else logo_paragraph.add_run()
            picture = run.add_picture(img_data, width=Inches(1.5))  # 1.5 pouces de largeur max
            
            return logo_paragraph
    except Exception as e:
        print(f"Erreur lors du téléchargement du logo: {e}")
        return None
    
    return None

def create_header_with_logo_and_title(doc, logo_url, title):
    """Créer l'en-tête avec titre à gauche et logo à droite"""
    # Créer un tableau invisible pour aligner titre (gauche) et logo (droite)
    header_table = doc.add_table(rows=1, cols=2)
    header_table.style = 'Table Grid'
    
    # Supprimer les bordures du tableau
    for row in header_table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.color.rgb = RGBColor(0, 0, 0)
    
    # Cellule de gauche : Titre
    title_cell = header_table.cell(0, 0)
    title_paragraph = title_cell.paragraphs[0]
    title_run = title_paragraph.add_run(title)
    title_run.font.size = Pt(20)
    title_run.bold = True
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    
    # Cellule de droite : Logo
    logo_cell = header_table.cell(0, 1)
    logo_paragraph = logo_cell.paragraphs[0]
    logo_paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Télécharger et ajouter le logo
    if logo_url:
        try:
            response = requests.get(logo_url, timeout=10)
            if response.status_code == 200:
                img_data = BytesIO(response.content)
                run = logo_paragraph.add_run()
                run.add_picture(img_data, width=Inches(1.2))  # 1.2 pouces de largeur
        except Exception as e:
            print(f"Erreur lors du téléchargement du logo: {e}")
    
    # Supprimer les bordures du tableau
    tbl = header_table._tbl
    for row in tbl.tr_lst:
        for cell in row.tc_lst:
            tcPr = cell.tcPr
            tcBorders = OxmlElement("w:tcBorders")
            for border in ['top', 'left', 'bottom', 'right']:
                border_elm = OxmlElement(f"w:{border}")
                border_elm.set(qn("w:val"), "nil")
                tcBorders.append(border_elm)
            tcPr.append(tcBorders)
    
    return header_table

def generate_docx_devis(devis, theme='bleu'):
    """Générer un DOCX de devis modifiable avec thème coloré et logo"""
    # Récupérer les couleurs du thème
    couleurs = THEMES_COULEURS_DOCX.get(theme, THEMES_COULEURS_DOCX['bleu'])
    
    filename = os.path.join('generated', f'devis_{devis.numero}_{theme}.docx')
    doc = Document()
    
    # Styles du document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # NOUVEAU: En-tête avec titre et logo
    create_header_with_logo_and_title(doc, devis.logo_url, "DEVIS")
    
    # Nom de l'entreprise avec couleur du thème
    company = doc.add_paragraph()
    company.add_run(devis.fournisseur_nom.upper()).bold = True
    company.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    company.runs[0].font.size = Pt(16)
    company.runs[0].font.color.rgb = couleurs['principale']
    
    doc.add_paragraph()  # Espace
    
    # Informations du devis
    info_table = doc.add_table(rows=3, cols=2)
    info_table.style = 'Light List'
    
    # Remplir le tableau d'informations
    info_data = [
        ('Numéro de devis:', devis.numero),
        ('Date d\'émission:', devis.date_emission),
        ('Date d\'expiration:', devis.date_expiration)
    ]
    
    for i, (label, value) in enumerate(info_data):
        info_table.cell(i, 0).text = label
        info_table.cell(i, 1).text = value
        # Mettre en gras les labels
        info_table.cell(i, 0).paragraphs[0].runs[0].bold = True
    
    doc.add_paragraph()  # Espace
    
    # Informations Fournisseur / Client
    doc.add_heading('ÉMETTEUR', level=2)
    doc.add_paragraph(f'{devis.fournisseur_nom}\n{devis.fournisseur_adresse}\n{devis.fournisseur_ville}')
    doc.add_paragraph(f'Email: {devis.fournisseur_email}\nTél: {devis.fournisseur_telephone}\nSIRET: {devis.fournisseur_siret}')
    
    doc.add_heading('CLIENT', level=2)
    doc.add_paragraph(f'{devis.client_nom}\n{devis.client_adresse}\n{devis.client_ville}')
    doc.add_paragraph(f'SIRET: {devis.client_siret}\nN° TVA: {devis.client_tva}')
    if devis.client_email:
        doc.add_paragraph(f'Email: {devis.client_email}')
    if devis.client_telephone:
        doc.add_paragraph(f'Tél: {devis.client_telephone}')
    
    doc.add_paragraph()  # Espace
    
    # Texte d'introduction si présent
    if devis.texte_intro:
        doc.add_paragraph(devis.texte_intro)
        doc.add_paragraph()
    
    # Tableau des articles
    items_table = doc.add_table(rows=1, cols=5)
    items_table.style = 'Table Grid'
    items_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    
    # En-têtes avec fond coloré selon le thème
    headers = ['Description', 'Qté', 'Prix unitaire', 'TVA (%)', 'Total HT']
    header_cells = items_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        # Mettre en gras et colorer avec le thème
        run = header_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)  # Blanc
        set_cell_background(header_cells[i], couleurs['header_bg'])
        # Alignement
        if i > 0:
            header_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Articles
    for item in devis.items:
        row = items_table.add_row()
        cells = row.cells
        
        # Description avec détails
        desc_text = item.description
        if item.details:
            desc_text += '\n' + '\n'.join([f'• {detail}' for detail in item.details])
        cells[0].text = desc_text
        
        # Données numériques
        cells[1].text = str(item.quantite)
        cells[2].text = f'{item.prix_unitaire:.2f} €'
        cells[3].text = f'{item.tva_taux} %'
        cells[4].text = f'{item.total_ht:.2f} €'
        
        # Alignement des cellules numériques
        for i in range(1, 5):
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Remise si applicable
        if item.remise > 0:
            remise_row = items_table.add_row()
            remise_cells = remise_row.cells
            remise_cells[3].text = 'Remise'
            remise_cells[4].text = f'-{item.remise:.2f} €'
            remise_cells[4].paragraphs[0].runs[0].font.color.rgb = RGBColor(231, 76, 60)  # Rouge
            remise_cells[3].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            remise_cells[4].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    doc.add_paragraph()  # Espace
    
    # Totaux
    totals_table = doc.add_table(rows=3, cols=2)
    totals_table.style = 'Light List'
    
    totals_data = [
        ('Total HT', f'{devis.total_ht:.2f} €'),
        ('TVA (20%)', f'{devis.total_tva:.2f} €'),
        ('TOTAL TTC', f'{devis.total_ttc:.2f} €')
    ]
    
    for i, (label, value) in enumerate(totals_data):
        totals_table.cell(i, 0).text = label
        totals_table.cell(i, 1).text = value
        totals_table.cell(i, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        totals_table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Mettre en gras et colorer la ligne TOTAL TTC avec couleur du thème
        if i == 2:
            for j in range(2):
                run = totals_table.cell(i, j).paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = couleurs['principale']
    
    doc.add_paragraph()  # Espace
    
    # Conditions de paiement
    doc.add_heading('CONDITIONS DE PAIEMENT', level=2)
    doc.add_paragraph(devis.conditions_paiement)
    doc.add_paragraph(devis.penalites_retard).runs[0].font.size = Pt(8)
    
    doc.add_paragraph()  # Espace
    
    # Informations bancaires
    doc.add_heading('COORDONNÉES BANCAIRES', level=2)
    bank_table = doc.add_table(rows=3, cols=2)
    bank_table.style = 'Light Grid'
    
    bank_data = [
        ('Banque:', devis.banque_nom),
        ('IBAN:', devis.banque_iban),
        ('BIC:', devis.banque_bic)
    ]
    
    for i, (label, value) in enumerate(bank_data):
        bank_table.cell(i, 0).text = label
        bank_table.cell(i, 1).text = value
        bank_table.cell(i, 0).paragraphs[0].runs[0].bold = True
    
    # Texte de conclusion
    if devis.texte_conclusion:
        doc.add_paragraph()
        doc.add_paragraph(devis.texte_conclusion)
    
    # Signature
    doc.add_paragraph()
    doc.add_paragraph()
    doc.add_paragraph('Bon pour accord')
    doc.add_paragraph('Date et signature:')
    doc.add_paragraph('_______________________')
    
    # Sauvegarder
    doc.save(filename)
    return filename

def generate_docx_facture(facture, theme='bleu'):
    """Générer un DOCX de facture modifiable avec thème coloré et logo"""
    # Récupérer les couleurs du thème
    couleurs = THEMES_COULEURS_DOCX.get(theme, THEMES_COULEURS_DOCX['bleu'])
    
    filename = os.path.join('generated', f'facture_{facture.numero}_{theme}.docx')
    doc = Document()
    
    # Styles du document
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)
    
    # NOUVEAU: En-tête avec titre et logo
    create_header_with_logo_and_title(doc, facture.logo_url, "FACTURE")
    
    # Nom de l'entreprise avec couleur du thème
    company = doc.add_paragraph()
    company.add_run(facture.fournisseur_nom.upper()).bold = True
    company.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    company.runs[0].font.size = Pt(16)
    company.runs[0].font.color.rgb = couleurs['principale']
    
    doc.add_paragraph()  # Espace
    
    # Informations de la facture
    info_table = doc.add_table(rows=6, cols=2)
    info_table.style = 'Light List'
    
    # Données de la facture
    info_data = [
        ('Numéro de facture:', facture.numero),
        ('Date d\'émission:', facture.date_emission),
        ('Date d\'échéance:', facture.date_echeance),
        ('Statut:', facture.statut_paiement),
        ('N° de commande:', facture.numero_commande),
        ('Réf. devis:', facture.reference_devis)
    ]
    
    row_index = 0
    for label, value in info_data:
        if value:  # N'ajouter que si la valeur existe
            info_table.cell(row_index, 0).text = label
            info_table.cell(row_index, 1).text = value
            info_table.cell(row_index, 0).paragraphs[0].runs[0].bold = True
            
            # Colorer le statut selon sa valeur
            if label == 'Statut:':
                run = info_table.cell(row_index, 1).paragraphs[0].runs[0]
                if value == 'En retard':
                    run.font.color.rgb = RGBColor(231, 76, 60)  # Rouge
                elif value == 'Payée':
                    run.font.color.rgb = RGBColor(39, 174, 96)  # Vert
                else:
                    run.font.color.rgb = couleurs['principale']  # Couleur du thème
                run.bold = True
            
            row_index += 1
    
    # Supprimer les lignes vides
    for i in range(row_index, 6):
        info_table._element.remove(info_table.rows[row_index]._element)
    
    doc.add_paragraph()  # Espace
    
    # Informations Fournisseur / Client
    doc.add_heading('ÉMETTEUR', level=2)
    doc.add_paragraph(f'{facture.fournisseur_nom}\n{facture.fournisseur_adresse}\n{facture.fournisseur_ville}')
    doc.add_paragraph(f'Email: {facture.fournisseur_email}\nTél: {facture.fournisseur_telephone}\nSIRET: {facture.fournisseur_siret}')
    
    doc.add_heading('CLIENT', level=2)
    doc.add_paragraph(f'{facture.client_nom}\n{facture.client_adresse}\n{facture.client_ville}')
    doc.add_paragraph(f'SIRET: {facture.client_siret}\nN° TVA: {facture.client_tva}')
    
    # Articles (même structure que devis)
    items_table = doc.add_table(rows=1, cols=5)
    items_table.style = 'Table Grid'
    
    # En-têtes avec couleur du thème
    headers = ['Description', 'Qté', 'Prix unitaire', 'TVA (%)', 'Total HT']
    header_cells = items_table.rows[0].cells
    for i, header in enumerate(headers):
        header_cells[i].text = header
        run = header_cells[i].paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(255, 255, 255)
        set_cell_background(header_cells[i], couleurs['header_bg'])
    
    # Ajouter les articles
    for item in facture.items:
        row = items_table.add_row()
        cells = row.cells
        
        desc_text = item.description
        if item.details:
            desc_text += '\n' + '\n'.join([f'• {detail}' for detail in item.details])
        
        cells[0].text = desc_text
        cells[1].text = str(item.quantite)
        cells[2].text = f'{item.prix_unitaire:.2f} €'
        cells[3].text = f'{item.tva_taux} %'
        cells[4].text = f'{item.total_ht:.2f} €'
        
        for i in range(1, 5):
            cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Totaux
    doc.add_paragraph()
    totals_table = doc.add_table(rows=3, cols=2)
    totals_table.style = 'Light List'
    
    totals_data = [
        ('Total HT', f'{facture.total_ht:.2f} €'),
        ('TVA (20%)', f'{facture.total_tva:.2f} €'),
        ('TOTAL TTC', f'{facture.total_ttc:.2f} €')
    ]
    
    for i, (label, value) in enumerate(totals_data):
        totals_table.cell(i, 0).text = label
        totals_table.cell(i, 1).text = value
        totals_table.cell(i, 0).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        totals_table.cell(i, 1).paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        if i == 2:  # Total TTC avec couleur du thème
            for j in range(2):
                run = totals_table.cell(i, j).paragraphs[0].runs[0]
                run.bold = True
                run.font.size = Pt(12)
                run.font.color.rgb = couleurs['principale']
    
    # Conditions et informations bancaires
    doc.add_paragraph()
    doc.add_heading('CONDITIONS DE PAIEMENT', level=2)
    doc.add_paragraph(facture.conditions_paiement)
    doc.add_paragraph(facture.penalites_retard).runs[0].font.size = Pt(8)
    
    doc.add_paragraph()
    doc.add_heading('COORDONNÉES BANCAIRES', level=2)
    bank_table = doc.add_table(rows=3, cols=2)
    bank_table.style = 'Light Grid'
    
    bank_data = [
        ('Banque:', facture.banque_nom),
        ('IBAN:', facture.banque_iban),
        ('BIC:', facture.banque_bic)
    ]
    
    for i, (label, value) in enumerate(bank_data):
        bank_table.cell(i, 0).text = label
        bank_table.cell(i, 1).text = value
        bank_table.cell(i, 0).paragraphs[0].runs[0].bold = True
    
    # Mentions légales
    doc.add_paragraph()
    doc.add_paragraph()
    legal = doc.add_paragraph()
    legal.add_run('Mentions légales: ').bold = True
    legal.add_run('TVA sur les encaissements. En cas de retard de paiement, seront exigibles, conformément à l\'article L441-10 du code de commerce, une indemnité calculée sur la base de trois fois le taux de l\'intérêt légal en vigueur ainsi qu\'une indemnité forfaitaire pour frais de recouvrement de 40 euros.')
    legal.runs[1].font.size = Pt(8)
    
    # Sauvegarder
    doc.save(filename)
    return filename
